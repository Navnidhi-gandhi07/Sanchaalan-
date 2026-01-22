import os
import re
import io
import sys
import json
from typing import List, Dict, Any, Tuple
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# --- Document I/O & Parsing ---
import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import pytesseract
from langdetect import detect

# Optional: Camelot for high-quality table extraction from PDFs
try:
    import camelot  # type: ignore
    _HAS_CAMELOT = True
except Exception:
    _HAS_CAMELOT = False

# --- NLP / Summarisation ---
from transformers import pipeline

# --- Office Docs ---
import docx
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx import Document

# --- Env ---
from dotenv import load_dotenv
load_dotenv()

# --- Mail ---
import smtp  # smtp.py in your project root

# =============================================================================
# Configuration
# =============================================================================
OUTPUT_DIR = "snapshots_out"
ASSETS_DIR = os.path.join(OUTPUT_DIR, "assets")
STRUCTURED_DIR = os.path.join(OUTPUT_DIR, "structured")
INDEX_DIR = os.path.join(OUTPUT_DIR, "index")

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(ASSETS_DIR, exist_ok=True)
os.makedirs(STRUCTURED_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)

# Stakeholder routing (editable)
STAKEHOLDER_EMAILS: Dict[str, List[str]] = {
    "Engineering": ["navnidhi.0717@gmail.com"],
    "Finance": ["navnidhi.0717@gmail.com"],
    "Safety": ["meghalgandhi.13@gmail.com"],
    "HR": ["meghalgandhi.13@gmail.com"],
    "Management": ["navnidhi.0717@gmail.com"],
}

# Role profiles
ROLE_PROFILES: Dict[str, Dict[str, Any]] = {
    "Engineering": {
        "focus": ["specification", "standard", "drawing", "maintenance", "downtime", "throughput", "SOP", "RCA", "root cause", "inspection"],
        "prompt": "Summarize for engineering: emphasize technical details, procedures, standards, deviations, and required actions.",
        "table_priority": ["spec", "part", "bom", "schedule", "inspection"],
    },
    "Finance": {
        "focus": ["cost", "price", "budget", "invoice", "capex", "opex", "amount", "tax", "gst", "penalty", "saving", "revenue"],
        "prompt": "Summarize for finance: extract monetary amounts, variances, due dates, cost drivers, and approvals required.",
        "table_priority": ["financial", "cost", "price", "budget", "invoice", "amount"],
    },
    "Safety": {
        "focus": ["incident", "hazard", "unsafe", "near miss", "ppe", "lockout", "danger", "risk", "compliance", "violation", "emergency"],
        "prompt": "Summarize for safety: highlight incidents, risks, mitigations, compliance obligations, and deadlines.",
        "table_priority": ["incident", "risk", "mitigation", "checklist"],
    },
    "HR": {
        "focus": ["policy", "leave", "holiday", "recruitment", "training", "benefit", "disciplinary", "overtime", "attendance", "payroll"],
        "prompt": "Summarize for HR: focus on policies, people actions, trainings, timelines, and required communications.",
        "table_priority": ["policy", "roster", "training", "attendance"],
    },
    "Management": {
        "focus": ["summary", "risk", "impact", "timeline", "decision", "escalation", "okrs", "kpi", "milestone"],
        "prompt": "Summarize for top management: provide an executive summary with key risks, decisions, timelines, and owners.",
        "table_priority": ["kpi", "milestone", "timeline", "budget"],
    },
}

# Summarizer (note: may download model first time if not cached)
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")


# =============================================================================
# Utilities
# =============================================================================
def detect_language(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return "unknown"
    try:
        return detect(text)
    except Exception:
        return "unknown"


def norm_spaces(text: str) -> str:
    text = re.sub(r"-\s+\n", "", text)
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, max_chars: int = 3500) -> List[str]:
    paras = [p.strip() for p in (text or "").split("\n") if p.strip()]
    chunks, cur = [], ""
    for p in paras:
        if len(cur) + len(p) + 1 <= max_chars:
            cur += (("\n" if cur else "") + p)
        else:
            if cur:
                chunks.append(cur)
            cur = p
    if cur:
        chunks.append(cur)
    return chunks


def ocr_pymupdf_page(page) -> str:
    pix = page.get_pixmap()
    mode = "RGB" if pix.alpha == 0 else "RGBA"
    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
    txt = pytesseract.image_to_string(img, lang="eng")
    return norm_spaces(txt)


def save_image_bytes(name_hint: str, image_bytes: bytes) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_.-]", "_", name_hint)
    path = os.path.join(ASSETS_DIR, safe)
    base, ext = os.path.splitext(path)
    i = 1
    final_path = path
    while os.path.exists(final_path):
        final_path = f"{base}_{i}{ext}"
        i += 1
    with open(final_path, "wb") as f:
        f.write(image_bytes)
    return final_path


def table_to_text(table_rows: List[List[str]]) -> str:
    if not table_rows:
        return ""
    out = []
    header = table_rows[0]
    for r in table_rows[1:6]:
        pairs = [f"{h.strip()}: {c.strip()}" for h, c in zip(header, r)]
        out.append("; ".join(pairs))
    return "\n".join(out)


# =============================================================================
# Extraction (Structured: text, tables, images)
# =============================================================================
def extract_pdf_structured(file_path: str) -> List[Dict[str, Any]]:
    results: List[Dict[str, Any]] = []
    doc = fitz.open(file_path)

    camelot_tables_by_page: Dict[int, List[List[List[str]]]] = {}
    if _HAS_CAMELOT:
        try:
            tables = camelot.io.read_pdf(file_path, pages="all", flavor="lattice")
            for t in tables:
                page = t.page if isinstance(t.page, int) else 0
                camelot_tables_by_page.setdefault(page, []).append(t.df.values.tolist())
        except Exception:
            camelot_tables_by_page = {}

    for i in range(len(doc)):
        page_idx = i + 1
        page = doc[i]

        raw = page.get_text("rawdict")  # type: ignore
        lines: List[str] = []
        for block in raw.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    spans = line.get("spans", [])
                    txt = "".join([s.get("text", "") for s in spans if "text" in s])
                    if txt.strip():
                        lines.append(txt)
        text = norm_spaces("\n".join(lines))

        if len(text) < 25:
            text = ocr_pymupdf_page(page)

        tables_list: List[Dict[str, Any]] = []
        if page_idx in camelot_tables_by_page:
            for tbl in camelot_tables_by_page[page_idx]:
                tables_list.append({"rows": [[str(c) for c in row] for row in tbl]})
        else:
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_pl = pdf.pages[i]
                    _tables = page_pl.extract_tables()
                    for t in _tables:
                        rows = [[(c or "").strip() for c in row] for row in t]
                        if any(any(cell for cell in row) for row in rows):
                            tables_list.append({"rows": rows})
            except Exception:
                pass

        images_meta: List[Dict[str, Any]] = []
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            try:
                base = doc.extract_image(xref)
                image_bytes = base.get("image", b"")
                ext = base.get("ext", "png")
                name_hint = f"{os.path.basename(file_path)}_p{page_idx}_img{img_index+1}.{ext}"
                path = save_image_bytes(name_hint, image_bytes)

                caption_text = ""
                try:
                    with Image.open(io.BytesIO(image_bytes)) as im:
                        caption_text = norm_spaces(pytesseract.image_to_string(im, lang="eng"))
                except Exception:
                    caption_text = ""
                if not caption_text:
                    caption_text = f"Figure on page {page_idx} (image {img_index+1})"

                images_meta.append({"path": path, "caption": caption_text})
            except Exception:
                continue

        lang = detect_language("\n".join([
            text,
            "\n".join([table_to_text(t.get("rows", [])) for t in tables_list]),
            "\n".join([im.get("caption", "") for im in images_meta])
        ]))

        results.append({
            "doc_id": os.path.basename(file_path),
            "page": page_idx,
            "type": "pdf",
            "lang": lang,
            "text": text,
            "tables": tables_list,
            "images": images_meta,
        })

    doc.close()
    return results


def extract_docx_structured(file_path: str) -> List[Dict[str, Any]]:
    d = docx.Document(file_path)
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    text = norm_spaces("\n".join(paras))
    return [{
        "doc_id": os.path.basename(file_path),
        "page": 1,
        "type": "docx",
        "lang": detect_language(text),
        "text": text,
        "tables": [],
        "images": [],
    }]


def extract_eml_structured(file_path: str) -> List[Dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        data = f.read()
    subject = re.search(r"^Subject:(.*)$", data, flags=re.MULTILINE)
    subject_text = subject.group(1).strip() if subject else ""
    body = re.split(r"\n\n", data, maxsplit=1)[-1]
    text = norm_spaces(f"Subject: {subject_text}\n\n{body}")
    return [{
        "doc_id": os.path.basename(file_path),
        "page": 1,
        "type": "email",
        "lang": detect_language(text),
        "text": text,
        "tables": [],
        "images": [],
    }]


def extract_any_structured(file_path: str) -> List[Dict[str, Any]]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_structured(file_path)
    if ext == ".docx":
        return extract_docx_structured(file_path)
    if ext == ".eml":
        return extract_eml_structured(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


# =============================================================================
# Hierarchical summarisation (map-reduce over long docs)
# =============================================================================
def summarize_long_text(text: str, max_chunk_chars: int = 3500) -> str:
    text = (text or "").strip()
    if not text:
        return "No content extracted."

    chunks = chunk_text(text, max_chunk_chars)
    partials: List[str] = []

    for ch in chunks:
        try:
            res = summarizer(ch, max_length=220, min_length=80, do_sample=False, truncation=True)
            partials.append(res[0]["summary_text"])  # type: ignore
        except Exception:
            partials.append(ch[:400])

    combined = "\n".join(partials)
    if len(combined) > max_chunk_chars:
        combined_chunks = chunk_text(combined, max_chunk_chars)
        reduced_parts = []
        for ch in combined_chunks:
            try:
                res = summarizer(ch, max_length=200, min_length=60, do_sample=False, truncation=True)
                reduced_parts.append(res[0]["summary_text"])  # type: ignore
            except Exception:
                reduced_parts.append(ch[:400])
        combined = "\n".join(reduced_parts)

    return combined


def build_role_context(structured_pages: List[Dict[str, Any]], role: str) -> str:
    profile = ROLE_PROFILES.get(role, {})
    focus_terms: List[str] = [t.lower() for t in profile.get("focus", [])]

    texts: List[str] = []
    for p in structured_pages:
        t = p.get("text", "")
        par_sel = []
        for line in t.split("\n"):
            line_l = line.lower()
            if any(ft in line_l for ft in focus_terms) or len(line) < 80 or re.match(r"^[A-Z0-9][A-Za-z0-9 .:/_-]{0,60}$", line):
                par_sel.append(line)
        if par_sel:
            texts.append("\n".join(par_sel))

    table_snippets: List[str] = []
    pri_tags = [t.lower() for t in profile.get("table_priority", [])]

    for p in structured_pages:
        for t in p.get("tables", []):
            rows = t.get("rows", [])
            header_join = ",".join([c.lower() for c in (rows[0] if rows else [])])
            score = sum(1 for tag in pri_tags if tag in header_join)
            if role == "Finance" or score > 0:
                table_snippets.append(table_to_text(rows))

    image_caps: List[str] = []
    if role in ("Engineering", "Safety", "Management"):
        for p in structured_pages:
            for im in p.get("images", []):
                cap = im.get("caption", "").strip()
                if cap:
                    image_caps.append(cap)

    parts = []
    parts.append(profile.get("prompt", f"Summarize for {role}."))

    if texts:
        parts.append("\n[ROLE-FOCUSED TEXT]\n" + "\n".join(texts[:50]))
    if table_snippets:
        parts.append("\n[IMPORTANT TABLE POINTS]\n" + "\n".join(table_snippets[:15]))
    if image_caps:
        parts.append("\n[FIGURE CAPTIONS]\n" + "\n".join(image_caps[:10]))

    return "\n\n".join(parts)


def role_specific_summary(structured_pages: List[Dict[str, Any]], generic_summary: str, role: str) -> str:
    role_context = build_role_context(structured_pages, role)
    combo = (role_context + "\n\n[GENERIC SUMMARY]\n" + (generic_summary or "")).strip()
    return summarize_long_text(combo)


# =============================================================================
# Indexing + Retrieval for Query Layer (LOCAL)
# =============================================================================
def build_chunks(structured_pages: List[Dict[str, Any]], max_chars: int = 900) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    for p in structured_pages:
        doc_id = p.get("doc_id")
        page = p.get("page")

        text = p.get("text", "")
        for i, ch in enumerate(chunk_text(text, max_chars=max_chars), start=1):
            chunks.append({
                "doc_id": doc_id,
                "page": page,
                "chunk_id": f"{doc_id}_p{page}_t{i}",
                "source": "text",
                "content": ch
            })

        for ti, t in enumerate(p.get("tables", []), start=1):
            table_txt = table_to_text(t.get("rows", []))
            if table_txt.strip():
                chunks.append({
                    "doc_id": doc_id,
                    "page": page,
                    "chunk_id": f"{doc_id}_p{page}_tbl{ti}",
                    "source": "table",
                    "content": table_txt
                })

        for ii, im in enumerate(p.get("images", []), start=1):
            cap = (im.get("caption") or "").strip()
            if cap:
                chunks.append({
                    "doc_id": doc_id,
                    "page": page,
                    "chunk_id": f"{doc_id}_p{page}_img{ii}",
                    "source": "image_caption",
                    "content": cap
                })
    return chunks


def _tokenize(s: str) -> List[str]:
    return re.findall(r"[a-zA-Z0-9]+", (s or "").lower())


def score_chunk(query: str, content: str) -> int:
    q_terms = _tokenize(query)
    c = (content or "").lower()
    score = 0
    for t in q_terms:
        if len(t) <= 2:
            continue
        score += c.count(t)
    return score


def retrieve_chunks(query: str, chunks: List[Dict[str, Any]], top_k: int = 6) -> List[Dict[str, Any]]:
    scored = []
    for ch in chunks:
        s = score_chunk(query, ch.get("content", ""))
        if s > 0:
            scored.append((s, ch))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [x[1] for x in scored[:top_k]]


def answer_query(query: str, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
    hits = retrieve_chunks(query, chunks, top_k=6)

    if not hits:
        return {
            "answer": "I couldn’t find relevant evidence in the uploaded documents.",
            "citations": []
        }

    evidence = "\n\n".join(
        [f"[CITE {i+1}] (Doc {h['doc_id']}, Page {h['page']}, {h['source']})\n{h['content']}"
         for i, h in enumerate(hits)]
    )

    prompt = (
        "Answer the question using ONLY the evidence below. "
        "If the evidence is insufficient, say you don't have enough information. "
        "Cite sources like [CITE 1], [CITE 2].\n\n"
        f"Question: {query}\n\nEvidence:\n{evidence}"
    )

    ans = summarize_long_text(prompt, max_chunk_chars=3500)

    citations = [
        {
            "cite": f"[CITE {i+1}]",
            "doc_id": h["doc_id"],
            "page": h["page"],
            "source": h["source"],
            "chunk_id": h["chunk_id"]
        }
        for i, h in enumerate(hits)
    ]
    return {"answer": ans, "citations": citations}


# =============================================================================
# Snapshot generation (DOCX)
# =============================================================================
def add_tables_to_docx(doc: DocxDocument, tables: List[List[List[str]]], max_tables: int = 4, max_rows: int = 12):
    count = 0
    for tbl in tables:
        if not tbl:
            continue
        header = tbl[0]
        rows = tbl[1:]
        doc.add_paragraph("")
        table = doc.add_table(rows=min(len(rows), max_rows) + 1, cols=len(header))
        table.style = "Light List"
        for j, h in enumerate(header):
            table.cell(0, j).text = str(h)
        for i, r in enumerate(rows[:max_rows], start=1):
            for j, c in enumerate(r):
                table.cell(i, j).text = str(c)
        count += 1
        if count >= max_tables:
            break


def add_images_to_docx(doc: DocxDocument, images_meta: List[Dict[str, Any]], max_images: int = 6):
    count = 0
    for im in images_meta:
        path = im.get("path")
        caption = im.get("caption", "")
        if not path or not os.path.exists(path):
            continue
        try:
            doc.add_paragraph("")
            doc.add_picture(path, width=Inches(5.5))
            if caption:
                p = doc.add_paragraph(caption)
                p.style = "Intense Quote"
            count += 1
            if count >= max_images:
                break
        except Exception:
            continue


def make_snapshot_docx(source_path: str, stakeholder: str, role_summary: str, structured_pages: List[Dict[str, Any]]):
    fname = f"{os.path.splitext(os.path.basename(source_path))[0]}__{stakeholder}_snapshot.docx"
    out_path = os.path.join(OUTPUT_DIR, fname)

    d = Document()
    d.add_heading(f"{stakeholder} Snapshot", 0)
    d.add_paragraph(f"Source: {os.path.basename(source_path)}")
    d.add_paragraph("")

    d.add_heading("Role-Specific Summary", level=1)
    for para in (role_summary or "No summary.").split("\n"):
        if para.strip():
            d.add_paragraph(para.strip())

    all_tables: List[List[List[str]]] = []
    all_images: List[Dict[str, Any]] = []
    for p in structured_pages:
        for t in p.get("tables", []):
            rows = t.get("rows", [])
            if rows:
                all_tables.append(rows)
        all_images.extend(p.get("images", []))

    if all_tables:
        d.add_paragraph("")
        d.add_heading("Key Tables", level=1)
        add_tables_to_docx(d, all_tables)

    if all_images and stakeholder in ("Engineering", "Safety", "Management"):
        d.add_paragraph("")
        d.add_heading("Figures & Diagrams", level=1)
        add_images_to_docx(d, all_images)

    d.add_paragraph("")
    d.add_paragraph(f"Traceability: {os.path.basename(source_path)}")
    d.save(out_path)
    return out_path


# =============================================================================
# Orchestration
# =============================================================================
def process_file(file_path: str) -> List[Tuple[str, str]]:
    """Main entry for app.py. Returns [(role, path_to_docx), ...]"""
    structured_pages = extract_any_structured(file_path)

    # save structured JSON
    struct_name = os.path.splitext(os.path.basename(file_path))[0] + "__structured.json"
    struct_path = os.path.join(STRUCTURED_DIR, struct_name)
    with open(struct_path, "w", encoding="utf-8") as f:
        json.dump(structured_pages, f, ensure_ascii=False, indent=2)

    # build + save index chunks (for Query Layer)
    chunks = build_chunks(structured_pages)
    idx_name = os.path.splitext(os.path.basename(file_path))[0] + ".chunks.json"
    idx_path = os.path.join(INDEX_DIR, idx_name)
    with open(idx_path, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    # build full text for generic summary
    text_blobs = []
    for p in structured_pages:
        text_blobs.append(p.get("text", ""))
        for t in p.get("tables", []):
            text_blobs.append(table_to_text(t.get("rows", [])))
        for im in p.get("images", []):
            text_blobs.append(im.get("caption", ""))
    full_text = norm_spaces("\n\n".join([t for t in text_blobs if t]))

    generic_summary = summarize_long_text(full_text)

    outputs: List[Tuple[str, str]] = []
    for role in ["Engineering", "Finance", "Safety", "HR", "Management"]:
        try:
            role_sum = role_specific_summary(structured_pages, generic_summary, role)
        except Exception:
            role_sum = generic_summary

        docx_path = make_snapshot_docx(file_path, role, role_sum, structured_pages)
        outputs.append((role, docx_path))

        # Email out (skip/fail-safe)
        try:
            smtp.send_email(
                subject=f"[Snapshot] {os.path.basename(file_path)} — {role}",
                body=(
                    "Attached is your role-specific snapshot.\n\n"
                    f"Structured JSON: {os.path.basename(struct_path)}\n"
                    f"Query Index: {os.path.basename(idx_path)}\n"
                ),
                recipients=STAKEHOLDER_EMAILS.get(role, []),
                attachments=[docx_path],
            )
        except Exception:
            pass

    return outputs


# =============================================================================
# CLI
# =============================================================================
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python mvp.py <file.pdf|.docx|.eml>")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f"File not found: {input_path}")
        sys.exit(1)

    outs = process_file(input_path)
    print("\n✅ Snapshots generated:")
    for role, path in outs:
        print(f" - {role}: {path}")

    # quick query demo
    print("\nTry a query:")
    q = input("Query: ").strip()
    if q:
        # load index we just wrote
        idx_name = os.path.splitext(os.path.basename(input_path))[0] + ".chunks.json"
        idx_path = os.path.join(INDEX_DIR, idx_name)
        with open(idx_path, "r", encoding="utf-8") as f:
            chs = json.load(f)
        res = answer_query(q, chs)
        print("\nAnswer:\n", res["answer"])
        print("\nCitations:\n", json.dumps(res["citations"], indent=2))
